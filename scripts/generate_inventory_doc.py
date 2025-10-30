from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_title(document: Document, text: str) -> None:
    title = document.add_heading(text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_paragraph(document: Document, text: str) -> None:
    document.add_paragraph(text)


def add_bullets(document: Document, items):
    for item in items:
        if isinstance(item, str):
            document.add_paragraph(item, style="List Bullet")
        elif isinstance(item, (list, tuple)) and len(item) > 0:
            # First element is parent, rest are children
            parent = item[0]
            document.add_paragraph(parent, style="List Bullet")
            for child in item[1:]:
                document.add_paragraph(child, style="List Bullet 2")
        elif isinstance(item, dict):
            # key: parent, value: list of children
            for k, v in item.items():
                document.add_paragraph(k, style="List Bullet")
                for child in v:
                    document.add_paragraph(child, style="List Bullet 2")


def add_numbered(document: Document, items):
    for item in items:
        document.add_paragraph(item, style="List Number")


def build_document() -> Document:
    document = Document()

    # Title
    add_title(document, "Complete Hardware Inventory System — End-to-End User Story Set")

    # Personas
    add_heading(document, "Personas", level=2)
    add_bullets(document, [
        "IT Asset Manager: Owns lifecycle policy, reporting, compliance.",
        "Procurement Officer: Manages vendors, POs, receiving.",
        "Technician: Tags, images, assigns, repairs, audits.",
        "Employee: Receives assets, acknowledges, requests service.",
        "Security/Compliance Auditor: Reviews trails, runs audits.",
        "Finance Partner: Tracks costs, depreciation, disposal value.",
        "Warehouse/Stock Clerk: Stores, picks, ships, reconciles stock.",
    ])

    # Scope and Assumptions
    add_heading(document, "Scope and Assumptions", level=2)
    add_bullets(document, [
        "Hardware types include laptops, desktops, servers, peripherals, network devices, and mobile devices.",
        "Integrations: SSO/IdP, HRIS (users/locations), ITSM (incidents/requests), MDM/Endpoint, Finance/ERP, Email/Slack.",
        "Assets uniquely identified via assetTag + serialNumber, with QR/Barcode/RFID support.",
    ])

    # Epic
    add_heading(document, "Epic: Manage Full Hardware Lifecycle From Request to Disposal", level=2)
    add_paragraph(document, (
        "As an IT Asset Manager, I need to control the entire lifecycle of hardware—"
        "from catalog and procurement through receiving, tagging, stock, assignment, support, audits, "
        "depreciation, and disposal—so that assets are tracked accurately, users are supported, and "
        "compliance/financial requirements are met."
    ))

    # End-to-End Narrative
    add_heading(document, "End-to-End Narrative (Happy Path)", level=2)
    add_numbered(document, [
        "Employee’s manager requests a laptop via portal (pre-approved bundle).",
        "Procurement creates a PO against an approved vendor and model; PO is approved.",
        "Shipment arrives; Receiving scans packing slip, records serials; status moves to Received.",
        "Technician tags devices, images them, records warranty and MDM enrollment; status In Stock.",
        "Auto-allocation triggers based on request; warehouse picks and ships to employee’s location.",
        "Employee acknowledges receipt; system records assignment; asset status In Use.",
        "During use, MDM feeds health; incidents create/link to maintenance tickets.",
        "Periodic audit runs; discrepancies resolved; audit trail captured.",
        "Finance runs monthly depreciation; reports published.",
        "At end-of-life, asset is wiped, de-assigned, and disposed via certified recycler; certificate and chain-of-custody stored; asset status Retired/Disposed.",
    ])

    # Acceptance Criteria by Phase
    add_heading(document, "Acceptance Criteria by Phase", level=2)

    add_heading(document, "Catalog & Vendors", level=3)
    add_bullets(document, [
        "Can define hardwareModel with specs, SKU, default warranty, cost.",
        "Maintain vendor list with terms, lead times, contacts, RMA process.",
        "Models can be active/inactive; blocking new POs when inactive.",
    ])

    add_heading(document, "Procurement", level=3)
    add_bullets(document, [
        "Create PO with vendor, line items, expected delivery, cost center.",
        "Approval workflow by amount and department.",
        "PO lifecycle: Draft → Pending Approval → Approved → Partially Received → Closed.",
        "Partial receipts supported; backorders tracked.",
    ])

    add_heading(document, "Receiving & Tagging", level=3)
    add_bullets(document, [
        "Receive by PO; scan serials to line items; record condition.",
        "Auto-generate unique assetTag; print QR/Barcode.",
        "Record location, warehouse bin, warranty start/end (by serial), invoice.",
        "Bulk receiving and import via CSV supported.",
    ])

    add_heading(document, "Warehousing & Stock", level=3)
    add_bullets(document, [
        "Real-time stock by location/bin; min/max thresholds.",
        "Transfer stock between locations with transfer orders and handoff logs.",
        "Cycle counts; variance resolution flow.",
    ])

    add_heading(document, "Imaging/Provisioning", level=3)
    add_bullets(document, [
        "Track imaging status, OS version, baseline software.",
        "MDM enrollment recorded; endpoint ID linked.",
        "Pre-assignment QA checklist stored.",
    ])

    add_heading(document, "Assignment / Check-Out", level=3)
    add_bullets(document, [
        "Assign to user or room; capture agreement/e-sign and handover photo.",
        "Shipment tracking until delivered; user acknowledgment required.",
        "Accessories bundle assignment (e.g., dock, charger) linked to main asset.",
    ])

    add_heading(document, "Moves / Check-In", level=3)
    add_bullets(document, [
        "Temporary check-out (loaners) with due dates and reminders.",
        "Check-in flow validates condition, accessories, and data wipe.",
    ])

    add_heading(document, "Maintenance / Support / RMA", level=3)
    add_bullets(document, [
        "Create maintenance ticket or link to existing ITSM ticket.",
        "Record diagnostics, parts used, vendor RMA, and downtime.",
        "Asset status transitions: In Use ↔ Under Repair ↔ In Stock with reasons.",
    ])

    add_heading(document, "Compliance & Audits", level=3)
    add_bullets(document, [
        "Scheduled audits by location or team; sampling and full counts supported.",
        "Scan assets; mark Found, Missing, Unexpected; variances resolved with notes.",
        "Immutable audit trail (who/when/where).",
    ])

    add_heading(document, "Finance & Depreciation", level=3)
    add_bullets(document, [
        "Asset cost basis from invoice; capex/opex flag; cost center.",
        "Straight-line depreciation schedule; monthly journal export to ERP.",
        "Net book value, impairment, write-offs, disposal proceeds tracked.",
    ])

    add_heading(document, "Disposition / Retirement", level=3)
    add_bullets(document, [
        "Enforce wipe/certification; attach certificates and chain-of-custody.",
        "Final valuation and proceeds; status Disposed.",
        "Prevent reassignment after retirement.",
    ])

    add_heading(document, "Reporting & Alerts", level=3)
    add_bullets(document, [
        "Dashboards: fleet composition, age, warranty coverage, assignments, audit compliance.",
        "Alerts: low stock, SLA breaches, expiring warranties, missing MDM heartbeat, overdue loans.",
        "Scheduled report delivery via email/Slack.",
    ])

    add_heading(document, "Security & Access", level=3)
    add_bullets(document, [
        "Roles: Admin, Asset Manager, Procurement, Technician, Auditor, Finance, Viewer.",
        "Field-level permissions for cost and PII.",
        "Full event audit log for all changes.",
    ])

    # Flow Overview
    add_heading(document, "Flow Overview", level=2)
    add_heading(document, "States", level=3)
    add_paragraph(document, (
        "Cataloged → On Order → Received → Tagged → In Stock → Allocated → In Transit → "
        "In Use → Under Repair → In Stock → Retired → Disposed"
    ))

    add_heading(document, "Key Transitions", level=3)
    add_bullets(document, [
        "PO Approved: Cataloged → On Order",
        "Receipt Posted: On Order → Received",
        "Tag + QA: Received → In Stock",
        "Assign/Ship: In Stock → In Transit → In Use",
        "Repair Created: In Use → Under Repair → In Stock/In Use",
        "End-of-Life: In Use/In Stock → Retired → Disposed",
    ])

    # Core Data Model
    add_heading(document, "Core Data Model (High-Level)", level=2)
    add_bullets(document, [
        "HardwareModel(id, name, sku, specs, defaultWarrantyMonths, active)",
        "Vendor(id, name, contacts, terms, rmaProcess)",
        "PurchaseOrder(id, vendorId, status, approvals[], lines[])",
        "Asset(id, assetTag, serialNumber, modelId, status, cost, invoiceId, warrantyStart, warrantyEnd, locationId, bin, mdmId, ownerUserId)",
        "Assignment(id, assetId, userId/locationId, startAt, endAt, acknowledgment)",
        "MaintenanceTicket(id, assetId, source, status, rmaId, parts[], downtimeHours)",
        "Audit(id, scope, startedAt, completedAt, findings[])",
        "Disposition(id, assetId, method, certificateUrl, proceeds, finalizedAt)",
        "Location(id, name, type, address)",
        "User(id, name, email, dept, managerId)",
        "Accessory(id, type, linkedAssetId, status)",
        "EventLog(id, entityType, entityId, actorId, action, timestamp, diff)",
    ])

    # Key Integrations
    add_heading(document, "Key Integrations", level=2)
    add_bullets(document, [
        "SSO/IdP: AuthN/Z and user provisioning.",
        "HRIS: Users, departments, locations; joiners/movers/leavers triggers.",
        "MDM/Endpoint: Enrollment, health, warranty, last seen.",
        "ITSM: Incidents/requests; ticket deep links.",
        "ERP/Finance: POs, invoices, depreciation journals, disposal proceeds.",
        "Notifications: Email/Slack for alerts and approvals.",
    ])

    # Non-Functional Requirements
    add_heading(document, "Non-Functional Requirements", level=2)
    add_bullets(document, [
        "Reliability: 99.9% uptime, eventual consistency for imports.",
        "Security: RBAC, audit logs, encrypted at rest/in transit, PII minimization.",
        "Performance: <300 ms p95 for primary views; bulk import 10k rows in <5 min.",
        "Scalability: Multi-location, 100k+ assets.",
        "Compliance: SOC2-style controls for auditability and access.",
    ])

    # Representative User Stories
    add_heading(document, "Representative User Stories (Selected)", level=2)

    add_heading(document, "Create and Maintain Catalog", level=3)
    add_bullets(document, [
        "As an Asset Manager, I can create and deactivate hardware models so only approved SKUs are ordered.",
        "Acceptance: Creating a model makes it selectable in POs; deactivating hides it from new POs.",
    ])

    add_heading(document, "Procure Hardware", level=3)
    add_bullets(document, [
        "As a Procurement Officer, I can submit POs for approval per spend policy and receive partial shipments.",
        "Acceptance: PO cannot close until all lines are fulfilled or canceled.",
    ])

    add_heading(document, "Receive and Tag", level=3)
    add_bullets(document, [
        "As a Technician, I can receive devices by PO, scan serials, auto-generate tags, and print labels.",
        "Acceptance: Duplicate serials are blocked; received counts update stock.",
    ])

    add_heading(document, "Assign to Employee", level=3)
    add_bullets(document, [
        "As a Technician, I can assign and ship assets; employees must acknowledge receipt.",
        "Acceptance: Status changes to In Use only after acknowledgment.",
    ])

    add_heading(document, "Maintain and Repair", level=3)
    add_bullets(document, [
        "As a Technician, I can move assets to Under Repair, record RMA, and return them to stock.",
        "Acceptance: While Under Repair, asset cannot be reassigned.",
    ])

    add_heading(document, "Audit Fleet", level=3)
    add_bullets(document, [
        "As an Auditor, I can run a site audit and reconcile variances with documented reasons.",
        "Acceptance: Audit closes only after all variances are resolved.",
    ])

    add_heading(document, "Depreciate Assets", level=3)
    add_bullets(document, [
        "As a Finance Partner, I can export monthly depreciation by cost center to ERP.",
        "Acceptance: Export is immutable and references asset IDs and journal lines.",
    ])

    add_heading(document, "Dispose Securely", level=3)
    add_bullets(document, [
        "As an Asset Manager, I can retire assets with wipe certificates and finalize disposal.",
        "Acceptance: Disposed assets cannot be reassigned; certificates are required.",
    ])

    # Definition of Done
    add_heading(document, "Definition of Done (System-Level)", level=2)
    add_bullets(document, [
        "All lifecycle states and transitions implemented with RBAC.",
        "Immutable event logs across procurement, receiving, assignment, maintenance, audit, finance, and disposal.",
        "Integrations live for SSO, HRIS sync, MDM heartbeat, and ERP exports.",
        "Dashboards and alerts configured; scheduled reports active.",
        "Data validation prevents duplicate serialNumber and assetTag.",
        (
            "End-to-end test proves the happy path and critical edge cases: Partial receipts, "
            "overdue loaner return, missing audit variance, replacement during RMA, employee offboarding reclaim."
        ),
    ])

    return document


if __name__ == "__main__":
    doc = build_document()
    output_path = "/workspace/Complete_Hardware_Inventory_System_User_Stories.docx"
    doc.save(output_path)
    print(output_path)
